# -*- coding: utf-8 -*-
"""สร้างไฟล์ Word คู่มือการใช้งาน Smart Bill Dashboard"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- Page setup ----------
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ---------- Default font (Thai-friendly) ----------
style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style.font.size = Pt(15)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:ascii'), 'TH Sarabun New')
rfonts.set(qn('w:hAnsi'), 'TH Sarabun New')
rfonts.set(qn('w:cs'), 'TH Sarabun New')
rfonts.set(qn('w:eastAsia'), 'TH Sarabun New')

# ---------- Helpers ----------
def add_title(text, size=24, color=(14, 116, 144)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(14, 116, 144)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0EA5E9')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(15, 23, 42)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(15)
    if bold:
        run.bold = True
    return p

def add_step(num, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {num}. ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(14, 165, 233)
    r1.font.size = Pt(15)
    r2 = p.add_run(text)
    r2.font.size = Pt(15)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(15)
    return p

def add_image_placeholder(caption):
    """กล่องว่างสำหรับ user แทรกรูป + caption บอกว่ารูปอะไร"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    # set border + bg
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    tcPr.append(shd)
    # border
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'dashed')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), '94A3B8')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # padding via empty paragraphs
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n[ แทรกภาพหน้าจอที่นี่ ]\n')
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 116, 139)
    cell.add_paragraph()  # spacer
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f'📷 {caption}')
    cap_run.italic = True
    cap_run.font.size = Pt(13)
    cap_run.font.color.rgb = RGBColor(71, 85, 105)
    cap.paragraph_format.space_after = Pt(10)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'💡 {text}')
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(180, 83, 9)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

# ============================================================
# Cover
# ============================================================
add_title('Smart Bill Dashboard')
add_title('คู่มือการใช้งาน', size=20, color=(71, 85, 105))
doc.add_paragraph().add_run('\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ระบบจัดการยอดขายร้านอาหาร / เครื่องดื่ม / กิจกรรมชายหาด')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph().add_run('\n\n')
add_image_placeholder('หน้าแรกของระบบ (Dashboard) — ภาพรวม')

doc.add_page_break()

# ============================================================
# 1. ภาพรวมระบบ
# ============================================================
add_h1('1. ภาพรวมระบบ')
add_para('Smart Bill Dashboard เป็นระบบบันทึกยอดขายและสรุปผลแบบเรียลไทม์ ใช้ Google Sheets เป็นฐานข้อมูล และเข้าถึงผ่านเว็บเบราว์เซอร์ได้จากทุกอุปกรณ์ (มือถือ / แท็บเล็ต / คอมพิวเตอร์)')
add_para('')
add_h2('เมนูหลัก 3 หน้า')
add_step(1, 'Dashboard — ดูภาพรวม / ยอดขาย / สถิติ / ตารางสรุป')
add_step(2, 'กรอกบิล — บันทึกการขายแต่ละรายการ')
add_step(3, 'จัดการเมนู — เพิ่ม / แก้ / ลบเมนูสินค้า')

add_image_placeholder('แถบเมนูหลัก 3 ปุ่มด้านบน (Dashboard / กรอกบิล / จัดการเมนู)')

doc.add_page_break()

# ============================================================
# 2. การเริ่มต้นใช้งาน (จัดการเมนู)
# ============================================================
add_h1('2. ตั้งค่าเริ่มต้น — เพิ่มเมนูสินค้า')
add_para('ก่อนเริ่มกรอกบิล ต้องเพิ่มเมนูสินค้าก่อน เช่น ลาเต้, กุ้งเผา, Sunbed, Jet Ski เป็นต้น')

add_h2('ขั้นตอนเพิ่มเมนูใหม่')
add_step(1, 'คลิกแท็บ "จัดการเมนู" บนแถบเมนูหลัก')
add_step(2, 'คลิกปุ่ม "+ เพิ่มเมนูใหม่" มุมขวาบน')
add_image_placeholder('หน้าจัดการเมนู — ปุ่ม + เพิ่มเมนูใหม่ มุมขวาบน')

add_step(3, 'กรอกข้อมูล: เลือก Emoji, ชื่อเมนู, หมวด')
add_step(4, 'กรอกตัวเลือก / ราคา')
add_image_placeholder('ฟอร์มเพิ่มเมนูใหม่ — กรอก Emoji, ชื่อ, หมวด, ตัวเลือก/ราคา')

add_h2('เมนูที่มีหลายตัวเลือก (เช่น ร้อน/เย็น)')
add_para('ถ้าเมนูเดียวกันมีหลายขนาดหรือตัวเลือก เช่น ลาเต้ ร้อน ฿80 / เย็น ฿100 ให้กดปุ่ม "+ เพิ่มตัวเลือก" เพิ่มแถวใหม่ แล้วกรอกแต่ละตัวเลือกพร้อมราคา')
add_image_placeholder('ฟอร์มเพิ่มเมนูที่มีหลาย variants — ลาเต้ ร้อน 80 / เย็น 100')

add_note('ระบบจะรวมเมนูชื่อเดียวกันที่มีหลายตัวเลือก ให้แสดงเป็นการ์ดเดียวที่มีปุ่ม mini เลือกได้ในหน้ากรอกบิลอัตโนมัติ')

add_h2('แก้ไข / ลบเมนู')
add_step(1, 'แต่ละแถวมีปุ่ม ✏️ (แก้ไข) และ 🗑 (ลบ)')
add_step(2, 'ใช้ Toggle สวิตช์ในคอลัมน์ "สถานะ" เพื่อเปิด/ปิดเมนูชั่วคราวโดยไม่ต้องลบ')
add_step(3, 'เลือกหลายรายการพร้อมกันด้วย checkbox แล้วใช้ "ลบ / เปิด / ปิด" จากแถบด้านบน')
add_image_placeholder('หน้าจัดการเมนู — แสดง toggle, ปุ่มแก้/ลบ และ bulk action bar')

doc.add_page_break()

# ============================================================
# 3. กรอกบิล
# ============================================================
add_h1('3. กรอกบิลขาย')
add_para('หน้านี้ใช้บันทึกยอดขายแต่ละบิล สามารถเลือกช่องทางลูกค้า เลือกวันที่ และเลือกเมนูใส่ตะกร้าได้')

add_h2('ขั้นตอนกรอกบิล')
add_step(1, 'คลิกแท็บ "กรอกบิล"')
add_step(2, 'เลือกวันที่ (ค่าเริ่มต้น = วันนี้)')
add_step(3, 'เลือกช่องทางลูกค้า: Walk-In / By Boats / Korean Tour / K. Speedboat')
add_image_placeholder('หน้ากรอกบิล — ปุ่มเลือกช่องทาง 4 สี + ช่องเลือกวันที่')

add_step(4, 'คลิกเมนูที่ต้องการเพิ่มเข้าตะกร้า — สามารถใช้ช่องค้นหาหรือเลือกจากหมวดได้')
add_step(5, 'ปรับจำนวนในตะกร้าด้วยปุ่ม + / − หรือลบรายการด้วยปุ่ม ✕')
add_image_placeholder('หน้ากรอกบิล — แสดงเมนู grid ด้านซ้าย + ตะกร้าด้านขวา')

add_step(6, 'ตรวจสอบยอดรวม แล้วคลิกปุ่ม "💾 บันทึกบิล"')
add_step(7, 'รอ popup สำเร็จ — ตะกร้าจะถูกล้างอัตโนมัติพร้อมรับบิลถัดไป')
add_image_placeholder('Popup บันทึกสำเร็จ — แสดง bill ID และยอดรวม')

add_h2('Tips การใช้งานเร็ว')
add_bullet('พิมพ์ชื่อเมนูในช่องค้นหาแล้วกด Enter — เพิ่มเมนูแรกที่ตรงเข้าตะกร้าทันที')
add_bullet('ส่วน "ขายดีตอนนี้" จะแสดงเมนูที่ถูกขายมากที่สุดใน 30 วัน — กดเพิ่มได้รวดเร็ว')
add_bullet('สีของปุ่มบันทึกจะเปลี่ยนตามช่องทางที่เลือก (เช่น Walk-In = แดง, By Boats = เขียว) — ป้องกันลืมตั้งช่องทาง')

doc.add_page_break()

# ============================================================
# 4. Dashboard
# ============================================================
add_h1('4. Dashboard — ดูสถิติ / ภาพรวม')
add_para('Dashboard มี 3 sub-tab: Overview, รายวัน, ตารางสรุป')

add_h2('การกรองข้อมูล')
add_step(1, 'เลือก "เดือน" จาก dropdown (ค่าเริ่มต้น = เดือนปัจจุบัน)')
add_step(2, 'หรือเลือก "ช่วงวันที่" (from → to) เพื่อดูเฉพาะช่วงที่ต้องการ')
add_step(3, 'เลือก "ช่องทาง" ลูกค้าด้วยปุ่ม pill ด้านล่าง')
add_step(4, 'คลิก "🔄 รีเฟรช" เพื่อโหลดข้อมูลใหม่จาก Google Sheets')
add_image_placeholder('Dashboard filter bar — month dropdown + date range + channel pills')

add_h2('Sub-tab: Overview')
add_para('แสดงภาพรวมยอดขายเดือนที่เลือก ประกอบด้วย:')
add_bullet('การ์ด "วันนี้" — ยอดขายแบ่งตาม 4 ช่องทาง')
add_bullet('Hero card ยอดเดือนนี้ vs เดือนก่อน')
add_bullet('KPI 4 ใบ — จำนวนบิล, เฉลี่ย/บิล, วันที่ขายดีสุด, เมนูขายดีสุด')
add_bullet('ผลงานตามช่องทาง 4 ช่อง')
add_bullet('กราฟยอดรายวัน + Donut สัดส่วน')
add_bullet('Top 10 เมนูขายดี + Insights อัตโนมัติ')
add_image_placeholder('Dashboard Overview — แสดง KPI cards และ channel cards')

add_h2('Sub-tab: รายวัน')
add_para('ตารางยอดขายแบ่งตามวันและช่องทาง — สลับดูได้ทีละช่องทางผ่านปุ่ม สรุปรวม / Walk-In / By Boats / Korean Tour / Korean Speedboat')
add_image_placeholder('Sub-tab รายวัน — ตารางยอดต่อวัน + ปุ่มสลับช่องทาง')

add_h2('Sub-tab: ตารางสรุป (Matrix)')
add_para('ตารางสรุปยอดแบ่งตามช่องทาง × หมวดสินค้า ทำให้เห็นว่าช่องทางไหนขายหมวดไหนได้มาก')
add_image_placeholder('Sub-tab ตารางสรุป — Matrix table แบ่งกลุ่ม Regular / Korean Tour / Korean Speedboat')

doc.add_page_break()

# ============================================================
# 5. การจัดการข้อมูลใน Google Sheet
# ============================================================
add_h1('5. การจัดการฐานข้อมูล (Google Sheet)')
add_para('ข้อมูลทั้งหมดเก็บใน Google Sheet ที่ถูก link กับระบบ สามารถเข้าไปดู/แก้ไขโดยตรงได้ แต่แนะนำให้แก้ผ่านระบบเว็บเพื่อกัน format เพี้ยน')

add_h2('Custom Menu ใน Google Sheet')
add_para('เปิด Google Sheet จะเห็นเมนู 🏝️ Smart Bill บนแถบเมนู ใช้สำหรับ:')
add_bullet('🛠️ Setup Database — สร้าง schema ใหม่ทั้งหมด (ใช้ครั้งแรก)')
add_bullet('🔄 Reset Menus — ใส่ seed เมนูตัวอย่าง 43 รายการ')
add_bullet('🗑️ Clear Menus — ลบเมนูทั้งหมด')
add_bullet('🧹 Clear Bills — ลบ Bills + BillItems ทั้งหมด')
add_bullet('ℹ️ ตรวจสอบโครงสร้าง — เช็คว่า schema ครบ')
add_image_placeholder('Custom menu 🏝️ Smart Bill ใน Google Sheet')

doc.add_page_break()

# ============================================================
# 6. คำถามที่พบบ่อย
# ============================================================
add_h1('6. คำถามที่พบบ่อย (FAQ)')

add_h2('Q: ใช้งานบนมือถือได้ไหม?')
add_para('ได้ครับ ระบบ Responsive รองรับทุกอุปกรณ์ — เปิด URL เดียวกันบนมือถือ/แท็บเล็ตได้เลย')

add_h2('Q: ทำไมเมนูใหม่ที่เพิ่มแล้วยังไม่ขึ้น?')
add_para('เปิด Dashboard แล้วกดปุ่ม "🔄 รีเฟรช" หรือ refresh เบราว์เซอร์ ระบบมี cache เพื่อความเร็ว')

add_h2('Q: ลบบิลที่บันทึกผิดได้ไหม?')
add_para('ปัจจุบันต้องเข้า Google Sheet → ลบแถวที่ผิดในชีต Bills และ BillItems ด้วยตนเอง')

add_h2('Q: ใครเข้าระบบได้บ้าง?')
add_para('ขึ้นกับการตั้งค่าตอน Deploy Web App: "Anyone" = ใครก็เข้าได้ผ่าน URL / "Anyone with Google account" = ต้อง login')

add_h2('Q: ข้อมูลปลอดภัยไหม?')
add_para('ข้อมูลทั้งหมดเก็บใน Google Sheet ของบัญชีคุณเอง Apps Script รันด้วยสิทธิ์เจ้าของ Sheet เท่านั้น ไม่มี server กลาง')

doc.add_paragraph().add_run('\n')
add_para('— จบคู่มือ —', bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# Save
# ============================================================
output_path = r'e:\Projects\Google Apps Script\Smart Bil Dashboard Foods\docs\คู่มือการใช้งาน_Smart_Bill_Dashboard.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
